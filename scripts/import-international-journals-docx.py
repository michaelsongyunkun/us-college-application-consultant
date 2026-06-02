from argparse import ArgumentParser
from pathlib import Path
import re

from docx import Document


ENTRY_PATTERN = re.compile(r"^(\d{3})\.\s+(.+)$")
FIELD_LABELS = {
    "url": "期刊地址：",
    "field": "期刊领域：",
    "type": "期刊类型：",
    "description": "期刊介绍：",
}


def default_source_path() -> Path:
    desktop = Path.home() / "Desktop"
    exact = desktop / "高中生可投国际期刊目录_200plus.docx"
    if exact.exists():
        return exact
    matches = sorted(desktop.glob("高中生可投国际期刊目录*.docx"))
    if not matches:
        raise SystemExit("No 高中生可投国际期刊目录*.docx file found on the Desktop.")
    return matches[0]


def clean_lines(document: Document) -> list[str]:
    return [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]


def primary_direction(field: str) -> str:
    return re.split(r"[；;]", field or "", maxsplit=1)[0].strip()


def index_database(journal_type: str) -> str:
    if "Scopus" in journal_type:
        return "Scopus"
    if "DOAJ" in journal_type:
        return "DOAJ OA"
    if "学生期刊" in journal_type:
        return "学生期刊"
    return journal_type.strip() or "待复核"


def parse_entries(lines: list[str]) -> tuple[list[str], list[dict[str, str]]]:
    header_lines: list[str] = []
    entries: list[dict[str, str]] = []
    current: dict[str, str] | None = None

    for line in lines:
        entry_match = ENTRY_PATTERN.match(line)
        if entry_match:
            if current:
                entries.append(current)
            current = {
                "id": entry_match.group(1),
                "name": entry_match.group(2).strip(),
            }
            continue

        if current is None:
            header_lines.append(line)
            continue

        for key, label in FIELD_LABELS.items():
            if line.startswith(label):
                current[key] = line[len(label) :].strip()
                break

    if current:
        entries.append(current)
    return header_lines, entries


def table_notes(document: Document) -> list[str]:
    notes: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in notes:
                    notes.append(text)
    return notes


def convert_docx_to_markdown(source: Path) -> str:
    document = Document(source)
    header_lines, entries = parse_entries(clean_lines(document))
    if not entries:
        raise SystemExit("No numbered journal entries were found in the DOCX.")

    title = header_lines[0] if header_lines else "高中生可投国际期刊目录"
    lines = [f"# {title}", ""]

    for note in table_notes(document):
        lines.extend([f"> {note}", ""])

    if len(header_lines) > 1:
        lines.extend(["## 使用说明", ""])
        for line in header_lines[1:]:
            if line == "使用说明":
                continue
            lines.append(f"- {line}")
        lines.append("")

    lines.extend(["## 国际期刊条目", ""])
    for entry in entries:
        field = entry.get("field", "")
        journal_type = entry.get("type", "")
        lines.extend(
            [
                f"### {entry['id']}. {entry['name']}",
                f"- **期刊地址**：{entry.get('url', '')}",
                f"- **论文方向**：{primary_direction(field)}",
                f"- **期刊领域**：{field}",
                f"- **检索库**：{index_database(journal_type)}",
                f"- **期刊类型**：{journal_type}",
                f"- **期刊介绍**：{entry.get('description', '')}",
                "",
            ]
        )

    return "\n".join(lines).strip() + "\n"


def main() -> None:
    parser = ArgumentParser(description="Convert the international journal DOCX into local RAG Markdown data.")
    parser.add_argument("source", nargs="?", type=Path, default=None)
    parser.add_argument("--output", type=Path, default=Path("data/international-journals.md"))
    args = parser.parse_args()

    source = args.source or default_source_path()
    markdown = convert_docx_to_markdown(source)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text(markdown, encoding="utf-8")
    print(f"Wrote {args.output} from {source}")


if __name__ == "__main__":
    main()
